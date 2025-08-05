# app/core/document_processor.py
import aiohttp
import asyncio
import logging
import time
import hashlib
import io
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import tempfile
import os

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document

# Try to import PyMuPDF, but handle gracefully if not available
try:
    import fitz  # PyMuPDF for better PDF handling
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# Try to import pypdf as alternative
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

from app.config.settings import settings
from app.models.response_models import DocumentChunk
from app.utils.text_processing import TextProcessor
from app.utils.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles document download, processing, and text extraction
    """
    
    def __init__(self):
        self.text_processor = TextProcessor()
        self.session: Optional[aiohttp.ClientSession] = None
    
    async def _get_session(self) -> aiohttp.ClientSession:
        """Get or create aiohttp session"""
        if not self.session or self.session.closed:
            timeout = aiohttp.ClientTimeout(total=settings.REQUEST_TIMEOUT)
            self.session = aiohttp.ClientSession(timeout=timeout)
        return self.session
    
    async def validate_document_url(self, url: str) -> bool:
        """
        Validate if document URL is accessible
        """
        try:
            session = await self._get_session()
            async with session.head(url) as response:
                if response.status == 200:
                    content_type = response.headers.get('content-type', '').lower()
                    content_length = int(response.headers.get('content-length', 0))
                    
                    # Check file size
                    if content_length > settings.MAX_FILE_SIZE:
                        logger.warning(f"File too large: {content_length} bytes")
                        return False
                    
                    # Check content type
                    supported_types = [
                        'application/pdf',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'application/msword'
                    ]
                    
                    return any(supported_type in content_type for supported_type in supported_types)
                
                return False
                
        except Exception as e:
            logger.error(f"URL validation failed: {e}")
            return False
    
    async def download_document(self, url: str) -> Tuple[bytes, str]:
        """
        Download document from URL
        
        Returns:
            Tuple of (document_content, file_extension)
        """
        session = await self._get_session()
        
        for attempt in range(settings.MAX_RETRIES):
            try:
                logger.info(f"Downloading document (attempt {attempt + 1}): {url}")
                
                async with session.get(url) as response:
                    if response.status == 200:
                        content = await response.read()
                        
                        # Determine file type from content-type or URL
                        content_type = response.headers.get('content-type', '').lower()
                        
                        if 'pdf' in content_type or url.lower().endswith('.pdf'):
                            file_ext = 'pdf'
                        elif 'word' in content_type or url.lower().endswith(('.docx', '.doc')):
                            file_ext = 'docx'
                        else:
                            # Try to detect from content
                            if content.startswith(b'%PDF'):
                                file_ext = 'pdf'
                            elif b'PK' in content[:10]:  # ZIP-based format (DOCX)
                                file_ext = 'docx'
                            else:
                                raise DocumentProcessingError("Unsupported file format")
                        
                        logger.info(f"Downloaded {len(content)} bytes, detected as {file_ext}")
                        return content, file_ext
                    
                    else:
                        raise DocumentProcessingError(f"Failed to download: HTTP {response.status}")
            
            except Exception as e:
                logger.warning(f"Download attempt {attempt + 1} failed: {e}")
                if attempt == settings.MAX_RETRIES - 1:
                    raise DocumentProcessingError(f"Failed to download after {settings.MAX_RETRIES} attempts: {e}")
                
                # Wait before retry
                await asyncio.sleep(2 ** attempt)
    
    def extract_text_from_pdf(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF content using multiple methods for robustness
        """
        text = ""
        metadata = {"pages": 0, "method": "unknown"}
        
        # Method 1: Try PyMuPDF (fitz) if available - generally better for complex PDFs
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                pages_text = []
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    pages_text.append(page_text)
                
                text = "\n\n".join(pages_text)
                metadata = {"pages": len(doc), "method": "pymupdf"}
                doc.close()
                
                # If we got good text, return it
                if len(text.strip()) > 100:
                    logger.info(f"Extracted text using PyMuPDF: {len(text)} characters")
                    return text, metadata
            
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: Try pypdf if available (alternative to PyMuPDF)
        if HAS_PYPDF and not text.strip():
            try:
                pdf_reader = pypdf.PdfReader(io.BytesIO(content))
                pages_text = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                
                text = "\n\n".join(pages_text)
                metadata = {"pages": len(pdf_reader.pages), "method": "pypdf"}
                
                if len(text.strip()) > 100:
                    logger.info(f"Extracted text using pypdf: {len(text)} characters")
                    return text, metadata
            
            except Exception as e:
                logger.warning(f"pypdf extraction failed: {e}")
        
        # Method 3: Try pdfplumber - good for tables and structured content
        if not text.strip():
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    pages_text = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages_text.append(page_text)
                    
                    text = "\n\n".join(pages_text)
                    metadata = {"pages": len(pdf.pages), "method": "pdfplumber"}
                    
                    if len(text.strip()) > 100:
                        logger.info(f"Extracted text using pdfplumber: {len(text)} characters")
                        return text, metadata
            
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 4: Fallback to PyPDF2
        if not text.strip():
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                pages_text = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                
                text = "\n\n".join(pages_text)
                metadata = {"pages": len(pdf_reader.pages), "method": "pypdf2"}
                
                logger.info(f"Extracted text using PyPDF2: {len(text)} characters")
                return text, metadata
            
            except Exception as e:
                logger.error(f"All PDF extraction methods failed: {e}")
                raise DocumentProcessingError("Failed to extract text from PDF")
        
        if not text.strip():
            raise DocumentProcessingError("No text could be extracted from PDF")
    
    def extract_text_from_docx(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX content
        """
        try:
            # Save content to temporary file (required for python-docx)
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(content)
                tmp_file.flush()
                
                try:
                    doc = Document(tmp_file.name)
                    
                    # Extract text from paragraphs
                    paragraphs = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text)
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                paragraphs.append(" | ".join(row_text))
                    
                    text = "\n\n".join(paragraphs)
                    metadata = {
                        "paragraphs": len(doc.paragraphs),
                        "tables": len(doc.tables),
                        "method": "python-docx"
                    }
                    
                    logger.info(f"Extracted text from DOCX: {len(text)} characters")
                    return text, metadata
                
                finally:
                    # Clean up temporary file
                    os.unlink(tmp_file.name)
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise DocumentProcessingError(f"Failed to extract text from DOCX: {e}")
    
    def create_chunks(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """
        Split text into manageable chunks for processing
        """
        # Clean and preprocess text
        cleaned_text = self.text_processor.clean_text(text)
        
        # Split into chunks
        chunks_data = self.text_processor.split_into_chunks(
            cleaned_text,
            chunk_size=settings.CHUNK_SIZE,
            overlap=settings.CHUNK_OVERLAP
        )
        
        # Create DocumentChunk objects
        chunks = []
        for i, chunk_data in enumerate(chunks_data):
            chunk = DocumentChunk(
                chunk_id=f"chunk_{i:04d}",
                text=chunk_data['text'],
                start_char=chunk_data['start'],
                end_char=chunk_data['end'],
                page_number=chunk_data.get('page_number'),
                metadata={
                    "chunk_index": i,
                    "word_count": len(chunk_data['text'].split()),
                    "char_count": len(chunk_data['text']),
                    **metadata
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    async def process_document_from_url(self, document_url: str, request_id: str) -> Dict[str, Any]:
        """
        Main method to process document from URL
        
        Returns:
            Dict containing processed chunks and metadata
        """
        start_time = time.time()
        
        try:
            # Step 1: Download document
            content, file_type = await self.download_document(document_url)
            
            # Step 2: Extract text based on file type
            if file_type == 'pdf':
                text, extraction_metadata = self.extract_text_from_pdf(content)
            elif file_type == 'docx':
                text, extraction_metadata = self.extract_text_from_docx(content)
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")
            
            # Step 3: Create chunks
            chunks = self.create_chunks(text, extraction_metadata)
            
            processing_time = time.time() - start_time
            
            # Prepare result
            result = {
                'chunks': chunks,
                'metadata': {
                    'total_characters': len(text),
                    'total_chunks': len(chunks),
                    'file_type': file_type,
                    'processing_time': round(processing_time, 2),
                    'extraction_method': extraction_metadata.get('method', 'unknown'),
                    **extraction_metadata
                }
            }
            
            logger.info(f"[{request_id}] Document processing completed: "
                       f"{len(chunks)} chunks, {len(text)} characters, "
                       f"{processing_time:.2f}s")
            
            return result
        
        except Exception as e:
            logger.error(f"[{request_id}] Document processing failed: {e}")
            raise DocumentProcessingError(f"Document processing failed: {e}")
    
    async def close(self):
        """Clean up resources"""
        if self.session and not self.session.closed:
            await self.session.close()
